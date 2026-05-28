"""DOCX report generation."""

import os
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.models import Run, Result
from core.exceptions import ReportError

def generate_docx_report(run: Run, results: list[Result], screenshots_dir: Path, output_path: Path) -> Path:
    """Generate a Word document report."""
    try:
        doc = Document()
        
        # Set margins to 2.5cm
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            
        # Default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # --- 1. Cover Section ---
        title = doc.add_heading("QA Test Run Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run_elem in title.runs:
            run_elem.font.color.rgb = RGBColor(0x00, 0x70, 0xF3)
            
        # Horizontal line trick
        p = doc.add_paragraph()
        p_border = p.add_run()
        p.paragraph_format.space_after = Pt(24)
        
        info_table = doc.add_table(rows=7, cols=2)
        info_data = [
            ("Test Name:", run.test_name),
            ("Run ID:", run.id),
            ("Status:", run.status.value.upper()),
            ("Started:", run.started_at or run.created_at),
            ("Duration:", f"{run.duration_seconds:.1f}s" if run.duration_seconds else "-"),
            ("Consultant:", run.consultant or "-"),
            ("Pod:", run.pod or "-")
        ]
        
        for i, (label, val) in enumerate(info_data):
            row = info_table.rows[i].cells
            row[0].text = label
            row[0].paragraphs[0].runs[0].font.bold = True
            row[1].text = val
            
            if label == "Status:":
                if val == "PASSED":
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
                elif val == "FAILED":
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
            elif label == "Run ID:":
                row[1].paragraphs[0].runs[0].font.name = 'Courier New'

        doc.add_paragraph()
        
        pass_rate = (run.passed_count / run.step_count * 100) if run.step_count else 0
        blocks_filled = int(pass_rate / 10)
        blocks_empty = 10 - blocks_filled
        bar = ("█" * blocks_filled) + ("░" * blocks_empty)
        doc.add_paragraph(f"{bar} {pass_rate:.0f}% passed")
        
        # --- 2. Summary Table ---
        doc.add_heading("Execution Summary", level=2)
        sum_table = doc.add_table(rows=2, cols=3)
        sum_table.style = 'Light Shading Accent 1'
        
        hdr = sum_table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Total", "Passed", "Failed"
        
        row = sum_table.rows[1].cells
        row[0].text = str(run.step_count)
        row[1].text = str(run.passed_count)
        row[2].text = str(run.failed_count)
        for c in row:
            c.paragraphs[0].runs[0].font.bold = True
            
        doc.add_page_break()
        
        # --- 3. Step Details ---
        doc.add_heading("Step Results", level=2)
        
        for res in results:
            # Heading
            p = doc.add_paragraph()
            r = p.add_run(f"Step {res.sequence} — {res.description or res.action.value}")
            r.font.bold = True
            if res.status == 'passed':
                r.font.color.rgb = RGBColor(0x22, 0xC5, 0x5E)
            elif res.status == 'failed':
                r.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                
            # Table
            t = doc.add_table(rows=2, cols=4)
            t.style = 'Table Grid'
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = "Action", "Selector", "Duration", "Status"
            for cell in h: cell.paragraphs[0].runs[0].font.bold = True
            
            v = t.rows[1].cells
            v[0].text = res.action.value
            v[1].text = res.selector or "-"
            v[2].text = f"{res.duration_ms}ms" if res.duration_ms else "-"
            v[3].text = res.status.upper()
            
            # Error
            if res.status == 'failed' and res.error_message:
                p = doc.add_paragraph()
                r = p.add_run(f"Error: {res.error_message}")
                r.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                r.italic = True
                
            # Screenshot
            if res.screenshot_path and os.path.exists(res.screenshot_path):
                size_kb = os.path.getsize(res.screenshot_path) / 1024
                if size_kb <= 800:
                    try:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        r.add_picture(res.screenshot_path, width=Cm(15))
                        p = doc.add_paragraph(f"Screenshot: Step {res.sequence}")
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.runs[0].font.size = Pt(9)
                        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
                    except Exception as e:
                        doc.add_paragraph(f"[Screenshot omitted: {e}]")
                else:
                    doc.add_paragraph(f"[Screenshot too large to embed: {size_kb:.1f}KB]")
            
            doc.add_paragraph("_" * 60) # Separator

        # --- 4. Footer ---
        for section in doc.sections:
            footer = section.footer
            p = footer.paragraphs[0]
            p.text = f"QA Platform\t\t{run.created_at[:10]}"
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
            
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        raise ReportError(f"Failed to generate DOCX report: {e}")
